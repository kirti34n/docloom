"""DOCX report renderer: Document IR to a native .docx via python-docx.

Theme tokens drive every color and font; rich text spans become formatted
runs, links become real hyperlinks, and cited spans get superscript
reference numbers resolved against doc.sources.
"""

from __future__ import annotations

import io
from pathlib import Path
from urllib.parse import urlsplit

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor
from docx.text.run import Run

from . import chart_svg, diagram_svg, raster
from ..ir import (
    Artifact,
    Block,
    BulletList,
    Callout,
    Chart,
    Code,
    Diagram,
    Divider,
    Document,
    Formula,
    Heading,
    Image,
    NumberedList,
    Paragraph,
    Quote,
    RichText,
    Sheet,
    Span,
    StatRow,
    Table,
    cited_ids,
    normalize_table,
    plain,
    report_blocks,
    source_numbers,
    spans,
)
from ..theme import Theme, hex_to_rgb
from . import RenderError

MONO_FONT = "Consolas"
HEADING_SIZES = {1: 20, 2: 16, 3: 13, 4: 12}
BULLET_STYLES = ("List Bullet", "List Bullet 2", "List Bullet 3")
CALLOUT_TOKEN = {"info": "primary", "success": "accent", "warning": "muted", "danger": "text"}
MAX_IMAGE_WIDTH = Inches(6)
# column-width clamp: no column collapses to unreadable, none hogs the page,
# so an 8+ column table still fits the frame instead of overflowing it
MIN_COL_WIDTH = Inches(0.6)
# Table Grid uses fixed layout (table.autofit = False), so this cap is a hard
# ceiling: an unbreakable token longer than MAX_COL_WIDTH can no longer widen
# its column to fit, and Word will overflow the cell rather than resize it.
# Left as documented behaviour rather than special-cased (LOW risk, narrow
# case -- a single unbreakable token wider than 2.5in -- and any fix trades
# it for a different failure, e.g. a table that no longer fits the frame).
MAX_COL_WIDTH = Inches(2.5)
# rasterize charts/diagrams at 2x the SVG's own width so the embedded picture
# stays crisp when Word scales it to MAX_IMAGE_WIDTH (and when it is printed)
CHART_RASTER_PX = chart_svg.DEFAULT_WIDTH * 2
# diagrams solve to a wider landscape canvas than a chart (target_aspect
# 2.0 by default) and carry more small text (node labels, sublabels, tags),
# so they get a higher raster width floor to stay legible at MAX_IMAGE_WIDTH
DIAGRAM_RASTER_PX = 1600
_SAFE_SCHEMES = {"http", "https", "mailto"}  # matches html.py


def _safe_href(url: str) -> str | None:
    try:
        scheme = urlsplit(url).scheme.lower()
    except ValueError:
        return None
    return url if scheme in _SAFE_SCHEMES else None


def _rgb(color: str) -> RGBColor:
    return RGBColor(*hex_to_rgb(color))


def _shade_cell(cell, color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(shd)


def _bottom_border(paragraph, color: str, sz: int = 6) -> None:
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color.lstrip("#"))
    pbdr.append(bottom)
    paragraph._p.get_or_add_pPr().append(pbdr)


def _add_hyperlink(paragraph, span: Span, theme: Theme, color: str | None = None) -> None:
    r_id = paragraph.part.relate_to(
        span.link, RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    if span.code:  # rFonts must precede b/i/color/u in rPr
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), MONO_FONT)
        fonts.set(qn("w:hAnsi"), MONO_FONT)
        rpr.append(fonts)
    if span.bold:
        rpr.append(OxmlElement("w:b"))
    if span.italic:
        rpr.append(OxmlElement("w:i"))
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), (color or theme.primary).lstrip("#"))
    rpr.append(color_el)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = span.text
    r.append(t)
    link.append(r)
    paragraph._p.append(link)


def _add_spans(
    paragraph, rt: RichText, theme: Theme, numbers: dict[str, int],
    link_color: str | None = None,
) -> None:
    for span in spans(rt):
        if span.link and _safe_href(span.link):
            _add_hyperlink(paragraph, span, theme, link_color)
        else:
            run = paragraph.add_run(span.text)
            if span.bold:
                run.bold = True
            if span.italic:
                run.italic = True
            if span.code:
                run.font.name = MONO_FONT
        if span.cite and span.cite in numbers:
            ref = paragraph.add_run(f"[{numbers[span.cite]}]")
            ref.font.superscript = True
            ref.font.color.rgb = _rgb(theme.primary)


def _setup_styles(docx_doc, theme: Theme) -> None:
    normal = docx_doc.styles["Normal"]
    normal.font.name = theme.font_body
    normal.font.size = Pt(11)
    normal.font.color.rgb = _rgb(theme.text)
    for level, size in HEADING_SIZES.items():
        style = docx_doc.styles[f"Heading {level}"]
        style.font.name = theme.font_heading
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = _rgb(theme.primary)


def _logo_block(docx_doc, doc: Document) -> None:
    """A right-aligned brand logo above the title, if the document carries a
    usable one. Failures are swallowed so a bad image never breaks the render."""
    logo = doc.logo
    if not (logo and logo.path and Path(logo.path).is_file()):
        return
    par = docx_doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    try:
        # 0.5in: the shared logo target height also used by typst (1.27cm)
        # and html (3rem = 48px @96dpi), so the brand mark is a consistent
        # size across every rendered format.
        par.add_run().add_picture(str(logo.path), height=Inches(0.5))
    except Exception:
        # remove the now-empty paragraph so no stray gap remains
        par._element.getparent().remove(par._element)


def _title_block(docx_doc, doc: Document, theme: Theme) -> None:
    _logo_block(docx_doc, doc)
    par = docx_doc.add_paragraph()
    run = par.add_run(doc.title)
    run.font.name = theme.font_heading
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = _rgb(theme.text)
    if doc.subtitle:
        run = docx_doc.add_paragraph().add_run(doc.subtitle)
        run.italic = True
        run.font.color.rgb = _rgb(theme.muted)
    byline = ", ".join(doc.authors)
    if doc.date:
        byline = f"{byline} \u2014 {doc.date}" if byline else doc.date
    if byline:
        run = docx_doc.add_paragraph().add_run(byline)
        run.font.size = Pt(9)
        run.font.color.rgb = _rgb(theme.muted)
    _bottom_border(docx_doc.add_paragraph(), theme.primary)


def _render_list(
    docx_doc, items, styles: tuple[str, ...], theme: Theme, numbers: dict[str, int]
) -> None:
    for item in items:
        style = styles[min(item.level, len(styles) - 1)]
        par = docx_doc.add_paragraph(style=style)
        _add_spans(par, item.text, theme, numbers)


def _render_numbered(docx_doc, items, theme: Theme, numbers: dict[str, int]) -> None:
    # ponytail: literal "1." prefix runs, so every list (and sub-level)
    # restarts at 1; the built-in List Number style shares one numbering
    # instance and keeps counting across lists. Upgrade path: a fresh
    # native <w:num> instance per list for real Word numbering.
    counters: dict[int, int] = {}
    for item in items:
        counters = {lv: n for lv, n in counters.items() if lv <= item.level}
        counters[item.level] = counters.get(item.level, 0) + 1
        par = docx_doc.add_paragraph()
        par.paragraph_format.left_indent = Inches(0.25 * (item.level + 1))
        par.add_run(f"{counters[item.level]}. ")
        _add_spans(par, item.text, theme, numbers)


def _render_quote(docx_doc, block: Quote, theme: Theme, numbers: dict[str, int]) -> None:
    par = docx_doc.add_paragraph()
    par.paragraph_format.left_indent = Inches(0.5)
    _add_spans(par, block.text, theme, numbers)
    # par.runs skips runs nested in w:hyperlink; walk the XML to catch them
    for r in par._p.iter(qn("w:r")):
        run = Run(r, par)
        run.italic = True
        run.font.color.rgb = _rgb(theme.muted)
    if block.attribution:
        par = docx_doc.add_paragraph()
        par.paragraph_format.left_indent = Inches(0.5)
        run = par.add_run("\u2014 " + block.attribution)
        run.font.size = Pt(10)
        run.font.color.rgb = _rgb(theme.muted)


def _render_code(docx_doc, block: Code, theme: Theme) -> None:
    table = docx_doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _shade_cell(cell, theme.surface)
    par = cell.paragraphs[0]
    for i, line in enumerate(block.code.split("\n")):
        if i:
            par = cell.add_paragraph()
        par.paragraph_format.space_after = Pt(0)
        run = par.add_run(line)
        run.font.name = MONO_FONT
        run.font.size = Pt(9.5)
    docx_doc.add_paragraph()


def _frame_width(docx_doc) -> Emu:
    """Usable page width (page width minus both margins) that a table may
    fill without overflowing onto -- or off of -- the page."""
    section = docx_doc.sections[0]
    return section.page_width - section.left_margin - section.right_margin


def _column_widths(docx_doc, header, rows, cols: int) -> list[Emu] | None:
    """Proportional column widths from a content-length proxy (the longest
    cell's plain text per column), capped at MAX_COL_WIDTH and renormalized
    so the columns still sum to the frame width -- capping alone can push
    the total under it, and a table that's still narrower than the frame
    is exactly the underflow this is meant to fix.

    MIN_COL_WIDTH is enforced AFTER that renormalization, not before: a
    naive clamp-then-renormalize (the previous approach) rescales every
    column -- including ones already pinned to the floor -- by the same
    factor, which silently pushes them back below it. Past ~10 columns
    that meant every column collapsed under the documented minimum while
    autofit was disabled, so Word rendered the collapsed widths verbatim
    as unreadable slivers instead of auto-sizing them away.

    Returns None if no split can honor MIN_COL_WIDTH for every column at
    all (i.e. cols * MIN_COL_WIDTH alone already exceeds the frame) -- the
    caller falls back to Word's own autofit for that table rather than
    emitting below-floor slivers."""
    if cols * MIN_COL_WIDTH > _frame_width(docx_doc):
        return None
    weights = []
    for j in range(cols):
        longest = len(plain(header[j]))
        for row in rows:
            longest = max(longest, len(plain(row[j])))
        weights.append(max(longest, 3))  # floor: an empty column still gets a share
    frame_width = _frame_width(docx_doc)
    total_weight = sum(weights)
    raw = [frame_width * w / total_weight for w in weights]
    capped = [min(w, MAX_COL_WIDTH) for w in raw]
    scale = frame_width / sum(capped)
    widths = [w * scale for w in capped]

    # Water-fill the floor: pin every below-floor column to MIN_COL_WIDTH
    # and take its shortfall proportionally from the columns still above
    # the floor. Pinning can push a donor below the floor in turn (it just
    # gave part of its width away), so repeat until nothing is left below
    # it. Each pass permanently pins at least one more column, so this
    # always terminates, and the guard above guarantees a feasible split
    # exists (the total width taken from donors never exceeds what they
    # have to give).
    pinned = [False] * cols
    while True:
        deficient = [i for i in range(cols) if not pinned[i] and widths[i] < MIN_COL_WIDTH]
        if not deficient:
            break
        shortfall = sum(MIN_COL_WIDTH - widths[i] for i in deficient)
        for i in deficient:
            pinned[i] = True
            widths[i] = MIN_COL_WIDTH
        donors = [i for i in range(cols) if not pinned[i]]
        donor_total = sum(widths[i] for i in donors)
        for i in donors:
            widths[i] -= shortfall * widths[i] / donor_total
    return [Emu(round(w)) for w in widths]


def _apply_column_widths(table, widths: list[Emu]) -> None:
    """python-docx requires the width set on every cell, not just the
    table.columns entry, for Word to actually honor non-uniform widths."""
    for column, width in zip(table.columns, widths):
        column.width = width
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width


def _repeat_header_row(table) -> None:
    """<w:tblHeader/> on the first row so it repeats on every page a long
    table spills onto, instead of the reader losing the column labels."""
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:tblHeader"))


def _prevent_row_splitting(table) -> None:
    """<w:cantSplit/> on every row so a single row never breaks across a
    page boundary mid-cell."""
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def _add_caption(
    docx_doc, caption: str, theme: Theme, alignment=None, keep_with_previous: bool = True,
) -> None:
    """A caption paragraph directly under a figure. keep_together so the
    caption itself never splits across a page break, and (unless the caller
    is handling it separately, as tables do) keep_with_next on the figure's
    own last paragraph so the figure and its caption are never torn apart
    across a page boundary -- one of the most recognizable "machine-
    generated" tells in the wild."""
    if keep_with_previous and docx_doc.paragraphs:
        docx_doc.paragraphs[-1].paragraph_format.keep_with_next = True
    par = docx_doc.add_paragraph()
    if alignment is not None:
        par.alignment = alignment
    run = par.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = _rgb(theme.muted)
    par.paragraph_format.keep_together = True


def _render_table(
    docx_doc, block: Table, theme: Theme, numbers: dict[str, int]
) -> None:
    header, rows = normalize_table(block.header, block.rows)
    cols = len(header)
    if not cols:  # fully empty table: zero-cell rows make Word offer "repair"
        return
    table = docx_doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = "Table Grid"
    widths = _column_widths(docx_doc, header, rows, cols)
    if widths is None:
        # Even splitting the frame width equally couldn't give every column
        # MIN_COL_WIDTH (too many columns for the page) -- there is no
        # explicit split that honors the documented floor, so let Word
        # autofit this table instead of emitting slivers thinner than it.
        table.autofit = True
    else:
        table.autofit = False  # otherwise Word ignores the explicit widths below
        _apply_column_widths(table, widths)
    _repeat_header_row(table)
    _prevent_row_splitting(table)
    for j, cell_rt in enumerate(header):
        cell = table.cell(0, j)
        _shade_cell(cell, theme.primary)
        par = cell.paragraphs[0]
        # header cells shade theme.primary, so a link there must recolor to
        # theme.background too or it renders primary-on-primary (invisible)
        _add_spans(par, cell_rt, theme, numbers, link_color=theme.background)
        for run in par.runs:
            run.bold = True
            run.font.color.rgb = _rgb(theme.background)
    for i, row in enumerate(rows):
        if i % 2 == 1:
            for j in range(cols):
                _shade_cell(table.cell(i + 1, j), theme.surface)
        for j, cell_rt in enumerate(row):
            _add_spans(table.cell(i + 1, j).paragraphs[0], cell_rt, theme, numbers)
    if block.caption:
        # the table itself isn't a paragraph docx_doc.paragraphs would see, so
        # keep_with_next goes on the last row's own cells instead, binding
        # the whole row (not just this new caption paragraph) to the page
        # the caption lands on
        for j in range(cols):
            table.cell(len(rows), j).paragraphs[-1].paragraph_format.keep_with_next = True
        _add_caption(docx_doc, block.caption, theme, keep_with_previous=False)
    else:
        docx_doc.add_paragraph()


def _render_image(docx_doc, block: Image, theme: Theme) -> bool:
    """Embed an image; returns True if a picture was actually added."""
    if not block.path:
        return False
    path = Path(block.path)
    if not path.is_file():
        return False
    try:
        picture = docx_doc.add_picture(str(path))
    except Exception:
        return False  # unembeddable (e.g. SVG) or unreadable
    if picture.width > MAX_IMAGE_WIDTH:
        picture.height = Emu(round(picture.height * MAX_IMAGE_WIDTH / picture.width))
        picture.width = MAX_IMAGE_WIDTH
    docx_doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if block.caption:
        _add_caption(docx_doc, block.caption, theme, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    return True


def _render_image_or_placeholder(docx_doc, block: Image, theme: Theme) -> None:
    """Like _render_image, but a file that exists and yet cannot be embedded
    (python-docx has no SVG decoder, so this is the common case for a
    diagram) leaves a labeled placeholder paragraph instead of nothing, so
    the figure never just vanishes. A genuinely missing path is still
    skipped silently, matching every other renderer."""
    if _render_image(docx_doc, block, theme):
        return
    if not block.path or not Path(block.path).is_file():
        return
    # python-docx cannot embed an SVG directly (so _render_image failed above),
    # but it can be rasterized -- the same path _render_chart/_render_diagram
    # take -- so an SVG Image/Artifact keeps its picture instead of degrading
    # to the "[image: alt]" text stub below.
    if raster.is_svg(block.path):
        png = raster.svg_file_to_png(
            block.path,
            width=DIAGRAM_RASTER_PX,
            font_files=raster.theme_font_files(theme),
        )
        if png is not None and _embed_png(docx_doc, png, block.caption, theme):
            return
    par = docx_doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run(f"[image: {block.alt}]" if block.alt else "[image]")
    run.italic = True
    run.font.color.rgb = _rgb(theme.muted)
    if block.caption:
        _add_caption(docx_doc, block.caption, theme, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def _embed_png(docx_doc, png: bytes, caption: str | None, theme: Theme) -> bool:
    """Add PNG bytes as a centered picture (plus caption). False if python-docx
    refuses the bytes, so the caller can still fall back."""
    par = docx_doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        par.add_run().add_picture(io.BytesIO(png), width=MAX_IMAGE_WIDTH)
    except Exception:
        return False
    if caption:
        _add_caption(docx_doc, caption, theme, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    return True


def _rasterize_chart_svg(svg: str, theme: Theme) -> bytes | None:
    """PNG bytes for a chart SVG, or None when the optional rasterizer extra
    (docloom[diagrams]) is not installed. None keeps the data-table fallback
    in _render_chart below, so a core install renders exactly as before."""
    return raster.svg_to_png(
        svg,
        width=CHART_RASTER_PX,
        font_files=raster.theme_font_files(theme),
    )


def _render_chart(
    docx_doc, block: Chart, theme: Theme, numbers: dict[str, int]
) -> None:
    # figure-style title: larger and on-brand, distinct from body paragraphs,
    # so the fallback below reads as a chart's data rather than a stray table
    if block.title:
        title_par = docx_doc.add_paragraph()
        run = title_par.add_run(block.title)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = _rgb(theme.primary)
        # keep the title glued to whatever renders next (the chart picture,
        # or the data-table fallback) so a page break can't strand the
        # title on one page and the chart on the next
        title_par.paragraph_format.keep_with_next = True
    png: bytes | None = None
    if block.path and Path(block.path).is_file():
        if _render_image(docx_doc, Image(path=block.path, caption=block.caption), theme):
            return
        # image present but unembeddable by python-docx: an SVG can still be
        # rasterized, anything else (corrupt file) falls through
        if raster.is_svg(block.path):
            png = raster.svg_file_to_png(
                block.path,
                width=CHART_RASTER_PX,
                font_files=raster.theme_font_files(theme),
            )
    if png is None:
        png = _rasterize_chart_svg(chart_svg.render_svg(block, theme), theme)
    if png is not None and _embed_png(docx_doc, png, block.caption, theme):
        return
    # no rasterizer available (docloom[diagrams] not installed): a titled,
    # captioned data table stands in for the chart instead of a bare,
    # unlabeled one.
    rows = [
        [s.name] + ["" if chart_svg._finite(v) is None else chart_svg._fmt(v) for v in s.values]
        for s in block.series
    ]
    _render_table(
        docx_doc,
        Table(header=[""] + list(block.labels), rows=rows, caption=block.caption),
        theme,
        numbers,
    )


def _diagram_theme(theme: Theme) -> dict:
    """diagram_svg's paint/solve pipeline takes a plain dict overlay, not the
    docloom Theme model (docs/diagram-plan.md section 3: "the docloom Theme
    model is adapted by callers"). Every renderer that embeds a diagram
    builds this same six-key adapter."""
    return {
        "primary": theme.primary,
        "accent": theme.accent,
        "surface": theme.surface,
        "text": theme.text,
        "muted": theme.muted,
        "background": theme.background,
    }


def _render_diagram(docx_doc, block: Diagram, theme: Theme) -> None:
    """Diagrams have no pre-rendered file (coordinate-free IR): the painter's
    SVG is the only source, and python-docx cannot embed SVG directly, so
    this always goes through raster.svg_to_png (unlike _render_chart, which
    can embed a pre-rendered picture path first). When the optional
    [diagrams] extra is absent (svg_to_png returns None), or the diagram
    itself is empty/malformed (solve() raises on anything lint would flag,
    e.g. a dangling edge), a labeled placeholder paragraph stands in instead
    of the figure just vanishing -- this function must never raise."""
    if not block.nodes:
        return
    try:
        svg = diagram_svg.render_svg(block, _diagram_theme(theme))
    except Exception:
        svg = ""
    png = (
        raster.svg_to_png(
            svg, width=DIAGRAM_RASTER_PX, font_files=raster.theme_font_files(theme)
        )
        if svg
        else None
    )
    if png is not None and _embed_png(docx_doc, png, block.caption, theme):
        return
    par = docx_doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run(f"[diagram: {block.alt}]" if block.alt else "[diagram]")
    run.italic = True
    run.font.color.rgb = _rgb(theme.muted)
    if block.caption:
        _add_caption(docx_doc, block.caption, theme, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def _render_stats(docx_doc, block: StatRow, theme: Theme) -> None:
    if not block.items:
        return
    table = docx_doc.add_table(rows=1, cols=len(block.items))
    for j, stat in enumerate(block.items):
        cell = table.cell(0, j)
        _shade_cell(cell, theme.surface)
        run = cell.paragraphs[0].add_run(stat.value)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = _rgb(theme.primary)
        run = cell.add_paragraph().add_run(stat.label)
        run.font.size = Pt(9)
        run.font.color.rgb = _rgb(theme.muted)
        if stat.delta:
            run = cell.add_paragraph().add_run(stat.delta)
            run.font.size = Pt(8)
            run.font.color.rgb = _rgb(theme.accent)
    docx_doc.add_paragraph()


def _render_callout(
    docx_doc, block: Callout, theme: Theme, numbers: dict[str, int]
) -> None:
    table = docx_doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _shade_cell(cell, theme.surface)
    par = cell.paragraphs[0]
    label = par.add_run(block.style.upper() + "  ")
    label.bold = True
    label.font.size = Pt(9)
    label.font.color.rgb = _rgb(getattr(theme, CALLOUT_TOKEN[block.style]))
    _add_spans(par, block.text, theme, numbers)
    docx_doc.add_paragraph()


def _render_block(
    docx_doc, block: Block, theme: Theme, numbers: dict[str, int]
) -> None:
    if isinstance(block, Heading):
        par = docx_doc.add_paragraph(style=f"Heading {min(block.level, 4)}")
        # never let a heading land alone at the bottom of a page, orphaned
        # from the body text it introduces
        par.paragraph_format.keep_with_next = True
        _add_spans(par, block.text, theme, numbers)
    elif isinstance(block, Paragraph):
        _add_spans(docx_doc.add_paragraph(), block.text, theme, numbers)
    elif isinstance(block, BulletList):
        _render_list(docx_doc, block.items, BULLET_STYLES, theme, numbers)
    elif isinstance(block, NumberedList):
        _render_numbered(docx_doc, block.items, theme, numbers)
    elif isinstance(block, Quote):
        _render_quote(docx_doc, block, theme, numbers)
    elif isinstance(block, Code):
        _render_code(docx_doc, block, theme)
    elif isinstance(block, Table):
        _render_table(docx_doc, block, theme, numbers)
    elif isinstance(block, Image):
        _render_image_or_placeholder(docx_doc, block, theme)
    elif isinstance(block, Callout):
        _render_callout(docx_doc, block, theme, numbers)
    elif isinstance(block, Chart):
        _render_chart(docx_doc, block, theme, numbers)
    elif isinstance(block, Diagram):
        _render_diagram(docx_doc, block, theme)
    elif isinstance(block, StatRow):
        _render_stats(docx_doc, block, theme)
    elif isinstance(block, Artifact):
        # picture embed if the render path exists; a labeled placeholder if
        # it exists but cannot be embedded (e.g. SVG); skip if there is none
        _render_image_or_placeholder(
            docx_doc, Image(path=block.path, alt=block.alt, caption=block.caption), theme
        )
    elif isinstance(block, Divider):
        _bottom_border(docx_doc.add_paragraph(), theme.muted)
    else:
        raise RenderError(f"unhandled block type: {type(block).__name__}")


def _sources_section(docx_doc, doc: Document, theme: Theme) -> None:
    numbers = source_numbers(doc)
    docx_doc.add_paragraph("Sources", style="Heading 1").paragraph_format.keep_with_next = True
    seen: set[str] = set()
    for source in doc.sources:
        if source.id in seen:  # duplicate id: numbers keeps the first, so skip the rest
            continue
        seen.add(source.id)
        par = docx_doc.add_paragraph()
        text = f"{numbers[source.id]}. {source.title}"
        if source.publisher:
            text += " \u2014 " + source.publisher
        if source.date:
            text += f" ({source.date})"
        if source.url:
            text += ", "
        run = par.add_run(text)
        run.font.size = Pt(10)
        if source.url:
            if _safe_href(source.url):
                _add_hyperlink(par, Span(text=source.url, link=source.url), theme)
            else:
                par.add_run(source.url).font.size = Pt(10)


def _sheet_cell_text(cell) -> str:
    if isinstance(cell, Formula):
        return cell.formula
    if cell is None:
        return ""
    if isinstance(cell, bool):
        return "TRUE" if cell else "FALSE"
    return str(cell)


def _render_sheet(docx_doc, sheet: Sheet, theme: Theme) -> None:
    heading = docx_doc.add_paragraph(style="Heading 2")
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run(sheet.name)
    run.bold = True
    _render_table(
        docx_doc,
        Table(
            header=[c.header for c in sheet.columns],
            rows=[[_sheet_cell_text(c) for c in row] for row in sheet.rows],
        ),
        theme,
        {},
    )


def render(doc: Document, theme: Theme, out_path: Path) -> Path:
    docx_doc = DocxDocument()
    _setup_styles(docx_doc, theme)
    _title_block(docx_doc, doc, theme)
    numbers = source_numbers(doc)
    for block in report_blocks(doc):
        _render_block(docx_doc, block, theme, numbers)
    for sheet in doc.sheets:  # workbooks would otherwise be silently dropped in DOCX
        _render_sheet(docx_doc, sheet, theme)
    if doc.sources and cited_ids(doc):
        _sources_section(docx_doc, doc, theme)
    docx_doc.save(str(out_path))
    return out_path
