"""Diagram-block wiring in the remaining renderers (docs/diagram-plan.md
section 4, phase P4): html inlines the solved SVG, typst embeds it as
native SVG bytes (true vector, no rasterizer), markdown writes a
diagram-{n}.svg sidecar, docx rasterizes it to a picture (or a placeholder
without the [diagrams] extra), and xlsx intentionally has no representation
for it (like Image). None of these renderers may raise, even on a diagram
solve() would reject (empty, dangling edge, ...); every one degrades to
either "skip silently" (nothing to draw) or a labeled placeholder (had
content, failed to render), matching the existing Image/Chart convention.
"""

from __future__ import annotations

import sys
import zipfile

import pytest

from docloom import (
    Diagram, DiagramEdge, DiagramGroup, DiagramNode, Document, Paragraph,
    Table as TableBlock, Theme, render,
)
from docloom.render import docx, html, markdown, typst


def _diagram(**overrides) -> Diagram:
    kwargs = dict(
        id="d1",
        title="Architecture",
        direction="LR",
        nodes=[
            DiagramNode(id="a", label="API", type="service", group="g1"),
            DiagramNode(id="b", label="DB", type="store", sublabel="Postgres 16"),
        ],
        edges=[DiagramEdge(source="a", target="b", label="writes")],
        groups=[DiagramGroup(id="g1", label="VPC")],
        caption="Figure 1: architecture",
        alt="architecture diagram",
    )
    kwargs.update(overrides)
    return Diagram(**kwargs)


def _dangling_diagram(**overrides) -> Diagram:
    """A diagram lint would reject (edge target not a node id): solve()
    raises a KeyError on this. Renderers must not propagate it."""
    kwargs = dict(
        nodes=[DiagramNode(id="a", label="A")],
        edges=[DiagramEdge(source="a", target="ghost")],
        caption="bad one",
        alt="bad diagram",
    )
    kwargs.update(overrides)
    return Diagram(**kwargs)


# --------------------------------------------------------------- html


def test_html_diagram_embeds_inline_svg_with_caption_and_aria_label():
    doc = Document(title="T", blocks=[_diagram()])
    out = html.to_html(doc, Theme())
    assert '<figure class="diagram"' in out
    assert "<svg" in out
    assert 'aria-label="architecture diagram"' in out
    assert "Figure 1: architecture" in out
    assert "<figcaption>" in out


def test_html_diagram_title_is_not_duplicated_outside_the_svg():
    # paint_svg already paints the diagram's own title inside the SVG; the
    # renderer must not also print a second, separate title element for it.
    doc = Document(title="T", blocks=[_diagram()])
    out = html.to_html(doc, Theme())
    assert out.count("Architecture") == 1


def test_html_diagram_with_no_nodes_renders_nothing():
    doc = Document(title="T", blocks=[_diagram(nodes=[], edges=[], groups=[])])
    out = html.to_html(doc, Theme())
    assert "<svg" not in out
    assert "<figure" not in out


def test_html_diagram_malformed_edge_does_not_raise_and_placeholders():
    # finding 14: html used to `except Exception: return ""` here, so a
    # dangling edge (solve() raises) vanished with zero trace. It must now
    # degrade to a visible placeholder and a runtime warning, matching
    # docx/markdown's existing "had content, failed to render" convention.
    doc = Document(title="T", blocks=[_dangling_diagram()])
    with pytest.warns(UserWarning, match="diagram"):
        out = html.to_html(doc, Theme())  # must not raise
    assert "<svg" not in out
    assert 'class="diagram-placeholder"' in out
    assert "bad diagram" in out  # the diagram's alt text, in the placeholder


def test_html_diagram_full_render_dispatch_via_top_level_api(tmp_path):
    doc = Document(title="T", blocks=[_diagram()])
    out = render(doc, "html", tmp_path / "d.html")
    text = out.read_text(encoding="utf-8")
    assert "<svg" in text


# --------------------------------------------------------------- typst


def test_typst_diagram_embeds_svg_bytes_with_caption():
    doc = Document(title="T", blocks=[_diagram()])
    out = typst.to_typst(doc, Theme())
    assert '#image(bytes(' in out
    assert 'format: "svg"' in out
    assert "Figure 1: architecture" in out


def test_typst_diagram_no_separate_title_text_block():
    # unlike Chart (whose SVG never paints a title), the diagram SVG already
    # draws d.title itself, so typst must not add its own #text(...) title
    doc = Document(title="T", blocks=[_diagram()])
    out = typst.to_typst(doc, Theme())
    assert '#text(weight: "bold")[Architecture]' not in out


def test_typst_diagram_no_nodes_produces_no_image_directive():
    doc = Document(title="T", blocks=[_diagram(nodes=[], edges=[], groups=[])])
    out = typst.to_typst(doc, Theme())
    assert "#image(bytes(" not in out


def test_typst_diagram_malformed_edge_does_not_raise():
    # finding 14: typst used to `except Exception: return []` here, so a
    # dangling edge (solve() raises) vanished with zero trace. It must now
    # degrade to a visible placeholder block plus a runtime warning, matching
    # docx/markdown's existing "had content, failed to render" convention.
    doc = Document(title="T", blocks=[_dangling_diagram()])
    with pytest.warns(UserWarning, match="diagram"):
        out = typst.to_typst(doc, Theme())  # must not raise
    assert "#image(bytes(" not in out
    assert "bad diagram" in out  # the diagram's alt text, in the placeholder


def test_typst_diagram_standalone_output_is_self_contained():
    # docs/diagram-plan.md section 4c and this module's own docstring promise
    # that to_typst() output works standalone, not only through render()'s
    # temp-dir compile path; the bytes-embed approach must not need any file
    # written next to a .typ that was never compiled.
    doc = Document(title="T", blocks=[_diagram()])
    out = typst.to_typst(doc, Theme())
    assert "docloom-image" not in out  # no unresolved external-file marker


# --------------------------------------------------------------- markdown


def test_markdown_diagram_writes_svg_sidecar_and_links_it(tmp_path):
    doc = Document(title="T", blocks=[_diagram()])
    out = markdown.render(doc, Theme(), tmp_path / "out.md")
    text = out.read_text(encoding="utf-8")
    assert "out_files/diagram-1.svg" in text
    assert "Figure 1: architecture" in text
    svg_path = tmp_path / "out_files" / "diagram-1.svg"
    assert svg_path.is_file()
    assert "<svg" in svg_path.read_text(encoding="utf-8")


def test_markdown_multiple_diagrams_get_sequential_filenames(tmp_path):
    doc = Document(
        title="T",
        blocks=[
            Paragraph(text="intro"),
            _diagram(id="d1", caption=None),
            _diagram(id="d2", caption=None),
        ],
    )
    out = markdown.render(doc, Theme(), tmp_path / "out.md")
    text = out.read_text(encoding="utf-8")
    assert "out_files/diagram-1.svg" in text
    assert "out_files/diagram-2.svg" in text
    assert (tmp_path / "out_files" / "diagram-1.svg").is_file()
    assert (tmp_path / "out_files" / "diagram-2.svg").is_file()


def test_markdown_diagram_no_nodes_skipped_silently(tmp_path):
    doc = Document(title="T", blocks=[_diagram(nodes=[], edges=[], groups=[], caption="empty")])
    out = markdown.render(doc, Theme(), tmp_path / "out.md")
    text = out.read_text(encoding="utf-8")
    assert text == "# T\n"
    assert not (tmp_path / "out_files").exists()


def test_markdown_diagram_malformed_gets_placeholder_not_raise(tmp_path):
    doc = Document(title="T", blocks=[_dangling_diagram()])
    out = markdown.render(doc, Theme(), tmp_path / "out.md")  # must not raise
    text = out.read_text(encoding="utf-8")
    assert "bad diagram" in text
    assert ".svg" not in text


def test_markdown_diagram_written_even_with_assets_false(tmp_path):
    # Diagrams are generated content, not a referenced file with an
    # "original location" to fall back on (unlike Image/Chart), so
    # assets=False must not silently drop them.
    doc = Document(title="T", blocks=[_diagram()])
    out = markdown.render(doc, Theme(), tmp_path / "noassets.md", assets=False)
    text = out.read_text(encoding="utf-8")
    assert "diagram-1.svg" in text
    assert (tmp_path / "noassets_files" / "diagram-1.svg").is_file()


# --------------------------------------------------------------- docx


def test_docx_diagram_embeds_a_picture(tmp_path):
    import docx as docx_lib

    doc = Document(title="T", blocks=[Paragraph(text="before"), _diagram(), Paragraph(text="after")])
    out = render(doc, "docx", tmp_path / "d.docx")
    d = docx_lib.Document(str(out))
    assert len(d.inline_shapes) == 1
    paragraphs = [p.text for p in d.paragraphs]
    assert "before" in paragraphs and "after" in paragraphs
    assert any("Figure 1: architecture" in p for p in paragraphs)
    with zipfile.ZipFile(out) as z:
        assert any(n.startswith("word/media/") for n in z.namelist())


def test_docx_diagram_without_rasterizer_extra_falls_back_to_placeholder(tmp_path, monkeypatch):
    import docx as docx_lib

    monkeypatch.setitem(sys.modules, "resvg_py", None)
    out = render(doc := Document(title="T", blocks=[_diagram()]), "docx", tmp_path / "d.docx")
    d = docx_lib.Document(str(out))
    assert len(d.inline_shapes) == 0
    text = "\n".join(p.text for p in d.paragraphs)
    assert "architecture diagram" in text
    assert "Figure 1: architecture" in text


def test_docx_diagram_malformed_edge_does_not_raise_and_placeholders(tmp_path):
    import docx as docx_lib

    doc = Document(title="T", blocks=[_dangling_diagram()])
    out = render(doc, "docx", tmp_path / "d.docx")  # must not raise
    d = docx_lib.Document(str(out))
    assert len(d.inline_shapes) == 0
    text = "\n".join(p.text for p in d.paragraphs)
    assert "bad diagram" in text


def test_docx_diagram_no_nodes_skipped_silently(tmp_path):
    import docx as docx_lib

    doc = Document(title="T", blocks=[
        Paragraph(text="only text here"),
        _diagram(nodes=[], edges=[], groups=[], alt="ghost diagram"),
    ])
    out = render(doc, "docx", tmp_path / "d.docx")
    d = docx_lib.Document(str(out))
    text = "\n".join(p.text for p in d.paragraphs)
    assert "ghost diagram" not in text
    assert len(d.inline_shapes) == 0


# --------------------------------------------------------------- xlsx


def test_xlsx_diagram_alongside_a_table_ignores_the_diagram(tmp_path):
    doc = Document(
        title="T",
        blocks=[TableBlock(header=["a"], rows=[["1"]]), _diagram()],
    )
    out = render(doc, "xlsx", tmp_path / "d.xlsx")
    assert out.is_file()
    with zipfile.ZipFile(out) as z:
        names = z.namelist()
        # exactly one data worksheet from the Table; no second sheet for the
        # diagram (it has no spreadsheet representation, like Image)
        sheet_xmls = [n for n in names if n.startswith("xl/worksheets/sheet")]
        assert len(sheet_xmls) == 1


def test_xlsx_diagram_only_document_raises_the_same_error_as_image_only(tmp_path):
    from docloom.render import RenderError

    doc = Document(title="T", blocks=[_diagram()])
    with pytest.raises(RenderError):
        render(doc, "xlsx", tmp_path / "d.xlsx")
