"""Re-audit regression: the DOCX chart data-table fallback must preserve
large-number fidelity. The old f"{v:g}" formatting capped at 6 significant
figures and flipped to scientific notation (1234567 -> "1.23457e+06"),
diverging from the SVG axis and silently changing the displayed value.

The data table is now the fallback taken only when the optional SVG rasterizer
(docloom[diagrams]) is absent, so the test poisons that import to reach it.
"""

import sys

from docloom import Chart, Document, Series, render


def test_docx_chart_table_keeps_large_numbers(tmp_path, monkeypatch):
    import docx as docx_lib

    monkeypatch.setitem(sys.modules, "resvg_py", None)  # force the table fallback
    doc = Document(
        title="Numbers",
        blocks=[
            Chart(
                title="Revenue",
                labels=["FY23", "FY24"],
                series=[Series(name="usd", values=[1234567.0, 1000000.0])],
            ),
        ],
    )
    d = docx_lib.Document(str(render(doc, "docx", tmp_path / "c.docx")))
    text = "\n".join(c.text for t in d.tables for row in t.rows for c in row.cells)
    assert "1,234,567" in text
    assert "1,000,000" in text
    assert "e+06" not in text
